from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime

doc = Document()

# ── Page margins ──────────────────────────────────────────────────────────────
for section in doc.sections:
    section.top_margin    = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin   = Inches(1.25)
    section.right_margin  = Inches(1.25)

# ── Helpers ───────────────────────────────────────────────────────────────────
def set_font(run, name="Calibri", size=11, bold=False, italic=False, color=None):
    run.font.name  = name
    run.font.size  = Pt(size)
    run.bold       = bold
    run.italic     = italic
    if color:
        run.font.color.rgb = RGBColor(*color)

def heading(text, level=1, color=(31, 73, 125)):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    sizes = {1: 16, 2: 13, 3: 11}
    set_font(run, size=sizes.get(level, 11), bold=True, color=color)
    p.paragraph_format.space_before = Pt(14 if level == 1 else 10)
    p.paragraph_format.space_after  = Pt(4)
    # bottom border for h1
    if level == 1:
        pPr = p._p.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')
        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'), 'single')
        bottom.set(qn('w:sz'), '6')
        bottom.set(qn('w:space'), '1')
        bottom.set(qn('w:color'), '1F497D')
        pBdr.append(bottom)
        pPr.append(pBdr)
    return p

def body(text, indent=0, space_after=4):
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_font(run, size=10.5)
    p.paragraph_format.left_indent  = Inches(indent * 0.3)
    p.paragraph_format.space_after  = Pt(space_after)
    return p

def bullet(text, level=0):
    p = doc.add_paragraph(style="List Bullet")
    run = p.add_run(text)
    set_font(run, size=10.5)
    p.paragraph_format.left_indent  = Inches(0.3 + level * 0.3)
    p.paragraph_format.space_after  = Pt(2)
    return p

def bold_label(para, label, rest, label_color=(31, 73, 125)):
    r1 = para.add_run(label)
    set_font(r1, bold=True, color=label_color)
    r2 = para.add_run(rest)
    set_font(r2)

def add_table(headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    # header row
    hdr = table.rows[0]
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(h)
        set_font(run, size=10, bold=True, color=(255, 255, 255))
        shading = OxmlElement('w:shd')
        shading.set(qn('w:val'), 'clear')
        shading.set(qn('w:color'), 'auto')
        shading.set(qn('w:fill'), '1F497D')
        cell._tc.get_or_add_tcPr().append(shading)
    # data rows
    for ri, row_data in enumerate(rows):
        row = table.rows[ri + 1]
        fill = 'EEF3FA' if ri % 2 == 0 else 'FFFFFF'
        for ci, val in enumerate(row_data):
            cell = row.cells[ci]
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            run = p.add_run(val)
            set_font(run, size=10)
            shading = OxmlElement('w:shd')
            shading.set(qn('w:val'), 'clear')
            shading.set(qn('w:color'), 'auto')
            shading.set(qn('w:fill'), fill)
            cell._tc.get_or_add_tcPr().append(shading)
    if col_widths:
        for ri, row in enumerate(table.rows):
            for ci, cell in enumerate(row.cells):
                cell.width = Inches(col_widths[ci])
    doc.add_paragraph()

def divider():
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(2)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '4')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), 'CCCCCC')
    pBdr.append(bottom)
    pPr.append(pBdr)

# ══════════════════════════════════════════════════════════════════════════════
#  COVER PAGE
# ══════════════════════════════════════════════════════════════════════════════
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(40)
r = p.add_run("SHIFTZY GO")
set_font(r, size=28, bold=True, color=(31, 73, 125))

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("India's Smart Vehicle Shifting & Travel Platform")
set_font(r, size=14, italic=True, color=(89, 89, 89))

doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("INSURANCE PARTNERSHIP PROPOSAL")
set_font(r, size=18, bold=True, color=(192, 80, 77))

doc.add_paragraph()
divider()
doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("A Comprehensive Proposal for Strategic Insurance Collaboration\n"
              "to Enable Safe, Legal, and Scalable Peer-to-Peer Vehicle Transportation")
set_font(r, size=11, italic=True, color=(89, 89, 89))

doc.add_paragraph()
doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run(f"Prepared: July 2026     |     Confidential & Proprietary")
set_font(r, size=10, color=(128, 128, 128))

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
#  1. EXECUTIVE SUMMARY
# ══════════════════════════════════════════════════════════════════════════════
heading("1. EXECUTIVE SUMMARY")
body(
    "Shiftzy Go is India's first peer-to-peer vehicle shifting and travel cost-sharing platform. "
    "We connect vehicle owners who need their car, bike, or SUV transported between cities with "
    "verified travelers heading to the same destination — allowing both parties to share the cost "
    "of fuel and tolls equally, making vehicle transport affordable and travel nearly free."
)
body(
    "Our platform creates a legally structured, technology-driven arrangement between two private "
    "individuals. For this arrangement to be legally sound, commercially scalable, and trusted by "
    "the public, a robust insurance framework is not optional — it is the single most critical "
    "pillar of our operations."
)
body(
    "This proposal outlines the exact insurance partnership we seek: a co-branded, API-driven, "
    "per-trip insurance product covering vehicle transit, third-party liability, and personal "
    "accident — embedded seamlessly into every Shiftzy Go booking. We are seeking an exclusive "
    "insurance partner who shares our vision of making vehicle transport safe, transparent, and "
    "affordable for every Indian."
)

# ══════════════════════════════════════════════════════════════════════════════
#  2. BUSINESS CONCEPT
# ══════════════════════════════════════════════════════════════════════════════
heading("2. THE BUSINESS CONCEPT — HOW SHIFTZY GO WORKS")

heading("2.1  The Problem We Solve", level=2)
body("Every day, millions of Indians face one of two common but expensive challenges:")
bullet("Vehicle owners need to relocate a car, bike, or SUV from one city to another — for job "
       "transfers, family relocations, vehicle purchases, or holiday returns. Current options are "
       "costly: truck/trailer carriers charge ₹8,000–₹30,000; hiring a professional driver is "
       "risky; driving themselves wastes time and money.")
bullet("Travelers pay the full cost of intercity travel — fuel, tolls, food, and time — often "
       "spending ₹2,500–₹10,000 for a single city-to-city journey.")

heading("2.2  The Shiftzy Go Solution", level=2)
body("Shiftzy Go matches both parties on the same route and date:")
bullet("The vehicle owner posts a 'shift request' listing the vehicle, pickup city, destination city, and preferred dates.")
bullet("Verified travelers browse open requests, apply to drive, and are approved after document verification.")
bullet("Both parties split the fuel and toll cost 50/50 through the app — a transparent, upfront calculation.")
bullet("A small platform fee plus GST is collected at checkout and displayed clearly before payment.")
bullet("The vehicle is delivered; the traveler reaches their destination — both at roughly half the usual cost.")

heading("2.3  The Math — A Real Example", level=2)
add_table(
    ["Scenario", "Cost"],
    [
        ["Full fuel + toll: Chennai to Bangalore (Sedan)", "₹4,200"],
        ["Vehicle owner's share (50%)", "₹2,100"],
        ["Traveler's share (50% fuel+toll + platform fee + GST)", "₹2,348"],
        ["Owner saves vs. truck carrier (avg. ₹9,000)", "₹6,900"],
        ["Traveler saves vs. bus/train + taxi", "₹1,500 – ₹2,000"],
    ],
    col_widths=[4.0, 2.0]
)

# ══════════════════════════════════════════════════════════════════════════════
#  3. WHY INSURANCE IS NON-NEGOTIABLE
# ══════════════════════════════════════════════════════════════════════════════
heading("3. WHY INSURANCE IS NON-NEGOTIABLE")
body(
    "Shiftzy Go's business model involves a third-party traveler legally driving a registered "
    "vehicle that belongs to someone else — across state borders, on public roads, for a "
    "financial consideration. This creates multiple layers of legal exposure:"
)
bullet("The vehicle owner's standard motor insurance policy covers only the registered owner "
       "and named drivers — not a third-party traveler. Any claim during a Shiftzy Go trip "
       "would likely be rejected under a standard policy.")
bullet("Under the Motor Vehicles Act, 1988, every vehicle on Indian roads must carry valid "
       "third-party liability insurance. When the traveler drives, the standard policy "
       "may not cover this — leaving both parties legally exposed.")
bullet("If the traveler is injured in an accident, no personal accident cover applies unless "
       "specifically arranged.")
bullet("Without insurance, neither party will trust the platform. No trust = no business.")
bullet("Regulators (IRDAI, state transport authorities) require clear insurance documentation "
       "for commercial-nature arrangements involving private vehicles.")

p = doc.add_paragraph()
r1 = p.add_run("Conclusion: ")
set_font(r1, bold=True, color=(192, 80, 77))
r2 = p.add_run(
    "An insurance partner who designs and issues a purpose-built per-trip insurance product "
    "for Shiftzy Go is not a vendor — they are a co-founder of this business model."
)
set_font(r2, size=10.5, italic=True)

# ══════════════════════════════════════════════════════════════════════════════
#  4. INSURANCE COVERAGE REQUIRED
# ══════════════════════════════════════════════════════════════════════════════
heading("4. INSURANCE COVERAGE WE REQUIRE FROM THE PARTNER")

heading("4.1  Per-Trip Vehicle Transit Insurance (Own Damage)", level=2)
body("This is the core product. It must be activated automatically on every confirmed booking.")
add_table(
    ["Parameter", "Details"],
    [
        ["Coverage", "Physical damage to the vehicle during the shift trip"],
        ["Who is insured", "The vehicle owner's registered vehicle (RC-based)"],
        ["Who drives", "Verified Shiftzy Go traveler (DL-verified, platform-approved)"],
        ["Policy trigger", "Booking confirmed on Shiftzy Go app"],
        ["Policy end", "Vehicle owner confirms delivery on app"],
        ["Sum insured", "IDV (Insured Declared Value) of the vehicle"],
        ["Key exclusions", "Pre-existing damage, drunk/drug-impaired driving, deliberate acts"],
        ["Documentation", "Pre-trip photos (time-stamped, geo-tagged), trip GPS log"],
    ],
    col_widths=[2.5, 4.0]
)

heading("4.2  Third-Party Liability (Mandatory — MV Act 1988)", level=2)
add_table(
    ["Parameter", "Details"],
    [
        ["Coverage", "Bodily injury, death, or property damage to third parties"],
        ["Minimum cover", "₹15 lakh per incident (as per MV Act, 1988)"],
        ["Applies to", "Traveler operating owner's vehicle during trip duration"],
        ["Legal basis", "Section 146, Motor Vehicles Act, 1988"],
    ],
    col_widths=[2.5, 4.0]
)

heading("4.3  Personal Accident Cover — Traveler", level=2)
add_table(
    ["Parameter", "Details"],
    [
        ["Coverage", "Accidental death or permanent total disability"],
        ["Sum assured", "₹10 lakh per traveler per trip (scalable)"],
        ["Who is covered", "The verified traveler approved for the shift"],
        ["Duration", "Trip start to delivery confirmation"],
    ],
    col_widths=[2.5, 4.0]
)

heading("4.4  Vehicle Owner Legal Liability Cover", level=2)
body(
    "Protects the vehicle owner from legal liability arising out of the traveler's actions "
    "during the shift — since the owner entrusted the vehicle to the traveler through our platform."
)

heading("4.5  Background Verification Support (Value-Added)", level=2)
body(
    "We request the insurance partner to provide or facilitate access to a KYC and driving "
    "history verification service or API:"
)
bullet("Driving License validity check (Parivahan database integration)")
bullet("Previous claim history of the traveler")
bullet("Identification of high-risk profiles before trip approval")
bullet("Aadhaar e-KYC for traveler identity confirmation")

# ══════════════════════════════════════════════════════════════════════════════
#  5. PROPOSED PARTNERSHIP MODEL
# ══════════════════════════════════════════════════════════════════════════════
heading("5. PROPOSED PARTNERSHIP MODEL")

heading("5.1  Embedded Micro-Insurance (Recommended Model)", level=2)
body(
    "The insurance premium is transparently built into the Shiftzy Go platform fee. Every "
    "confirmed trip automatically generates a short-term insurance policy — no separate "
    "purchase, no extra steps for the user. This is the global best practice for gig-economy "
    "insurance (Uber, Ola, Rapido all use embedded insurance)."
)
bullet("Shiftzy Go applies for IRDAI Corporate Agent (Composite) registration or works through a licensed insurance broker.")
bullet("Insurer designs a 'Shiftzy Go Vehicle Transit Policy' — a unique, co-branded product.")
bullet("Policy is issued via API at booking confirmation; certificate emailed to both parties instantly.")
bullet("Claims reported through Shiftzy Go app; insurer handles survey and settlement.")

heading("5.2  Alternative: Opt-In Insurance at Checkout", level=2)
body(
    "Insurance displayed as a separate optional line item at checkout. Users choose coverage tier:"
)
add_table(
    ["Tier", "Covers", "Estimated Premium"],
    [
        ["Basic", "Third-Party Liability only", "₹50 – ₹80 / trip"],
        ["Standard", "TP + Own Damage + PA (Traveler)", "₹150 – ₹300 / trip"],
        ["Comprehensive", "All covers + Owner Legal Liability", "₹300 – ₹600 / trip"],
    ],
    col_widths=[1.5, 3.5, 2.0]
)
body("We recommend the Embedded Model for maximum adoption and legal clarity.")

# ══════════════════════════════════════════════════════════════════════════════
#  6. REVENUE & PREMIUM PROJECTIONS
# ══════════════════════════════════════════════════════════════════════════════
heading("6. REVENUE & PREMIUM PROJECTIONS")

heading("6.1  Per-Trip Premium Estimates", level=2)
add_table(
    ["Vehicle Type", "Avg. Distance", "Suggested Premium/Trip", "Insurer Net (after 12% commission)"],
    [
        ["Bike / Two-Wheeler", "150 – 400 km", "₹80 – ₹150", "₹70 – ₹132"],
        ["Hatchback / Sedan", "200 – 600 km", "₹180 – ₹320", "₹158 – ₹282"],
        ["SUV / MUV", "200 – 600 km", "₹280 – ₹480", "₹246 – ₹422"],
        ["Luxury / Premium", "200 – 700 km", "₹450 – ₹800", "₹396 – ₹704"],
    ],
    col_widths=[1.8, 1.5, 1.8, 2.4]
)

heading("6.2  Volume & Revenue Forecast", level=2)
add_table(
    ["Year", "Target Cities", "Est. Monthly Trips", "Est. Annual Trips", "Est. Annual Premium to Insurer"],
    [
        ["Year 1", "7 cities (South India)", "500 – 800", "6,000 – 10,000", "₹15 – ₹25 lakh"],
        ["Year 2", "15 cities (Pan India)", "2,500 – 4,000", "30,000 – 48,000", "₹75 lakh – ₹1.2 crore"],
        ["Year 3", "30+ cities + corporate", "10,000+", "1,20,000+", "₹3 – ₹5 crore"],
    ],
    col_widths=[0.8, 2.0, 1.5, 1.5, 2.2]
)

heading("6.3  Commission Structure (Per IRDAI Norms)", level=2)
bullet("Shiftzy Go earns a distribution/referral commission of 10–15% of net premium collected")
bullet("Insurer receives 85–90% of premium net of applicable GST")
bullet("Volume-based bonuses may be negotiated as trip volumes scale")
bullet("No claims bonus (NCB) equivalent discount for routes/vehicles with zero-claim history")

# ══════════════════════════════════════════════════════════════════════════════
#  7. RISK MANAGEMENT
# ══════════════════════════════════════════════════════════════════════════════
heading("7. SHIFTZY GO'S RISK MANAGEMENT COMMITMENTS")
body(
    "We understand that a low claims ratio is in both our interests. Shiftzy Go commits to the "
    "following risk controls to protect the insurer and minimise fraudulent claims:"
)
add_table(
    ["Risk Control", "How It Works"],
    [
        ["Driving License Verification", "Every traveler must upload DL; verified against Parivahan/DigiLocker before first trip approval. Expired or suspended DLs are automatically rejected."],
        ["Aadhaar / KYC Verification", "Full Aadhaar e-KYC or PAN verification for every user before they can book or drive."],
        ["Pre-Trip Vehicle Photos", "Time-stamped, geo-tagged photos uploaded by both owner and traveler before trip start. Stored permanently on our servers for claim reference."],
        ["GPS Trip Tracking", "Real-time route tracking during transit. Alerts for route deviation, speeding, or prolonged stops. (Active in Phase 2)"],
        ["Rating & Blacklist System", "5-star rating system for travelers and owners. Users below 3.5 stars or with any fraud flag are suspended immediately."],
        ["Vehicle Condition Pre-screening", "Owners must upload RC, valid PUC certificate, and existing insurance copy before listing."],
        ["In-App Claim Reporting", "First Notice of Loss (FNOL) submitted within 24 hours via app with photos, GPS data, and trip details."],
        ["Fraud Detection", "Pattern analysis of claim frequency, route anomalies, and user behaviour. Suspected fraud escalated to insurer's SIU within 6 hours."],
    ],
    col_widths=[2.0, 4.5]
)

# ══════════════════════════════════════════════════════════════════════════════
#  8. TERMS & CONDITIONS OF THE PARTNERSHIP
# ══════════════════════════════════════════════════════════════════════════════
heading("8. TERMS & CONDITIONS OF THE PARTNERSHIP")

heading("8.1  Obligations of Shiftzy Go", level=2)
bullet("Shiftzy Go shall obtain and maintain IRDAI Corporate Agent registration (or operate through a licensed broker) before issuing any policies.")
bullet("Shiftzy Go shall remit collected premiums to the insurer within 7 working days of each calendar month-end.")
bullet("Shiftzy Go shall share trip data (GPS logs, pre-trip photos, user KYC details) with the insurer within 24 hours of any FNOL.")
bullet("Shiftzy Go shall not approve any traveler whose DL is expired, suspended, or unverified.")
bullet("Shiftzy Go shall maintain a dedicated Claims Liaison Officer as the single point of contact with the insurer.")
bullet("Shiftzy Go shall achieve a minimum trip volume of 500 insured trips/month within 6 months of commercial launch, failing which both parties shall renegotiate minimum guarantee terms.")
bullet("Shiftzy Go shall not partner with any other insurance company for the same product category during the exclusivity period (2 years) without written consent.")

heading("8.2  Obligations of the Insurance Partner", level=2)
bullet("The insurer shall design and file a 'Shiftzy Go Vehicle Transit Policy' under IRDAI's File & Use procedure within 90 days of this agreement being signed.")
bullet("The insurer shall provide API access for real-time policy issuance, status inquiry, and certificate delivery to both vehicle owner and traveler.")
bullet("The insurer shall acknowledge all FNOL claims within 48 hours of receipt.")
bullet("The insurer shall assign a licensed surveyor and complete vehicle inspection within 72 hours of claim acknowledgement.")
bullet("The insurer shall settle valid claims within 30 days of submission of complete documentation.")
bullet("The insurer shall provide a dedicated SPOC (Single Point of Contact) for platform integration, claims management, and escalations.")
bullet("The insurer shall offer a co-branded certificate of insurance for each trip (digital, via app and email).")
bullet("The insurer shall offer volume-based premium discounts as annual trip volumes cross thresholds of 12,000, 36,000, and 1,00,000 trips.")

heading("8.3  Exclusivity Terms", level=2)
bullet("The insurer is granted exclusive rights to provide vehicle transit micro-insurance on the Shiftzy Go platform for an initial period of 2 years from commercial launch.")
bullet("Exclusivity may be extended by mutual agreement in writing for further 1-year periods.")
bullet("Exclusivity lapses automatically if the insurer fails to meet claims SLA (settlement within 30 days) for more than 15% of claims in any rolling 6-month period.")

# ══════════════════════════════════════════════════════════════════════════════
#  9. LEGAL & REGULATORY COMPLIANCE
# ══════════════════════════════════════════════════════════════════════════════
heading("9. LEGAL & REGULATORY COMPLIANCE")
add_table(
    ["Regulation", "How Shiftzy Go Complies"],
    [
        ["Motor Vehicles Act, 1988", "All trips carry valid TP insurance; verified before trip start."],
        ["IRDAI (Insurance Regulatory and Development Authority)", "Shiftzy Go to register as Corporate Agent (Composite) or operate via licensed broker for policy distribution."],
        ["Insurance Act, 1938", "All premiums collected remitted within statutory timelines; proper books of account maintained."],
        ["Information Technology Act, 2000 & DPDP Act, 2023", "User data encrypted at rest and in transit; shared with insurer only with user consent; data retention policy maintained."],
        ["GST (Goods & Services Tax)", "18% GST on insurance premium collected separately and remitted to insurer for payment to government."],
        ["KYC / AML Norms (RBI & IRDAI)", "Full Aadhaar + DL verification for all travelers; CERSAI and CKYC integration planned."],
        ["Consumer Protection Act, 2019", "In-app grievance mechanism; escalation to IRDAI Bima Bharosa portal; response within 15 days."],
        ["Arbitration & Conciliation Act, 1996", "All disputes resolved through arbitration in Chennai with a sole arbitrator agreed upon by both parties."],
    ],
    col_widths=[2.2, 4.3]
)

# ══════════════════════════════════════════════════════════════════════════════
#  10. FUTURE PLANS
# ══════════════════════════════════════════════════════════════════════════════
heading("10. FUTURE PLANS — GROWTH ROADMAP")
add_table(
    ["Phase", "Timeline", "Geography", "Insurance Implication"],
    [
        ["Phase 1: South India Launch", "Months 1–12", "Chennai, Bangalore, Hyderabad, Coimbatore, Madurai, Tirunelveli, Kochi", "Basic transit + TP + PA; 500–800 trips/month target"],
        ["Phase 2: Pan-India Expansion", "Months 12–24", "Mumbai, Pune, Delhi NCR, Ahmedabad, Jaipur, Kolkata", "Premium tiers; corporate fleet accounts; bike + SUV surge"],
        ["Phase 3: Ecosystem Expansion", "Months 24–36", "All major metros + Tier 2 cities", "OEM delivery partnerships; annual subscription insurance for frequent users"],
        ["Phase 4: Market Leadership", "Years 3–5", "30+ cities; international corridors (India–Nepal, India–Sri Lanka pending permits)", "Co-branded insurance subsidiary; IPO readiness; 1 lakh+ trips/year"],
    ],
    col_widths=[1.5, 1.2, 2.3, 2.5]
)

# ══════════════════════════════════════════════════════════════════════════════
#  11. WHY PARTNER WITH SHIFTZY GO
# ══════════════════════════════════════════════════════════════════════════════
heading("11. WHY PARTNER WITH SHIFTZY GO?")
add_table(
    ["Factor", "Benefit to Insurance Partner"],
    [
        ["First-mover advantage", "Capture the peer-to-peer vehicle transit insurance market before any competitor enters. No existing product in India addresses this specific segment."],
        ["Digital-first, API-driven", "Zero paper. Real-time policy issuance. Automated certificate delivery. GPS trip data for claims. Significantly lower admin cost than traditional motor insurance."],
        ["Built-in risk controls", "Pre-trip photos, DL verification, GPS tracking, and rating systems reduce claim frequency and severity substantially."],
        ["High-growth premium potential", "₹15–25 lakh premium in Year 1 growing to ₹3–5 crore by Year 3 — a product line that scales automatically with the platform."],
        ["Young, urban, creditworthy customers", "Shiftzy Go users are primarily 21–40 year old working professionals — the most valuable insurance demographic in India."],
        ["Exclusive partnership offered", "We approach only one insurer for this product category. The partner gets 100% of all Shiftzy Go trip insurance volume."],
        ["Brand visibility", "Co-branded certificates, in-app insurer logo, and joint PR elevate the insurer's brand in the gig-economy and new-age insurance space."],
        ["Clean claims data", "Every trip has GPS data, photos, and user-verified documentation — making fraud detection and claims processing faster and cheaper."],
    ],
    col_widths=[2.0, 4.5]
)

# ══════════════════════════════════════════════════════════════════════════════
#  12. PROPOSED DEAL STRUCTURE
# ══════════════════════════════════════════════════════════════════════════════
heading("12. PROPOSED DEAL — TERM SHEET SUMMARY")

body("The following outlines the key commercial and legal terms of the proposed partnership agreement:")
doc.add_paragraph()

items = [
    ("Parties:", "Shiftzy Go Technologies Pvt. Ltd. (Platform) and [Insurance Company Name] (Insurer)"),
    ("Partnership Type:", "Embedded Insurance Distribution Agreement (Platform as Corporate Agent / Broker Arrangement)"),
    ("Agreement Duration:", "3 years from date of first policy issuance; auto-renewable on mutual written consent"),
    ("Product Name:", "Shiftzy Go Vehicle Transit Policy (co-branded)"),
    ("Policy Type:", "Short-term micro-insurance; per-trip basis; covers Own Damage + Third Party Liability + Personal Accident (Traveler)"),
    ("Premium Collection:", "Collected by Shiftzy Go at booking checkout; remitted to insurer within 7 working days of month-end"),
    ("Distribution Commission:", "12% of net premium to Shiftzy Go (or as per IRDAI norms); subject to annual renegotiation based on volumes"),
    ("Claims SLA:", "FNOL acknowledgement within 48 hours; survey within 72 hours; settlement within 30 days of complete documentation"),
    ("Exclusivity:", "Insurer holds exclusive rights in the Shiftzy Go peer-to-peer vehicle transit category for 2 years"),
    ("Minimum Volume Guarantee:", "500 insured trips/month from Month 6; 2,000/month from Month 18; 8,000/month from Month 30"),
    ("Pilot Programme:", "100-trip pilot in Chennai–Bangalore corridor before full commercial launch; claims reviewed jointly"),
    ("API Integration:", "Insurer to provide REST API for policy issuance, certificate delivery, and claim status within 60 days of agreement signing"),
    ("Governing Law:", "Laws of India"),
    ("Jurisdiction & Arbitration:", "Disputes submitted to sole arbitration in Chennai, Tamil Nadu per Arbitration & Conciliation Act, 1996"),
]
for label, value in items:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(3)
    r1 = p.add_run(f"  {label}  ")
    set_font(r1, bold=True, color=(31, 73, 125))
    r2 = p.add_run(value)
    set_font(r2, size=10.5)

# ══════════════════════════════════════════════════════════════════════════════
#  13. NEXT STEPS
# ══════════════════════════════════════════════════════════════════════════════
heading("13. NEXT STEPS")
steps = [
    "Share this proposal with your Product Innovation / New Business team",
    "Schedule a presentation meeting between Shiftzy Go founders and insurance leadership",
    "Sign a mutual Non-Disclosure Agreement (NDA) for detailed business data sharing",
    "Insurer's actuarial team reviews trip data and prices the per-trip product",
    "File & Use submission to IRDAI for the new product",
    "Technical team begins API integration (target: 60 days from agreement signing)",
    "Pilot: 100 insured trips on Chennai–Bangalore corridor; joint claims review",
    "Full commercial launch with embedded insurance on all bookings",
]
for i, s in enumerate(steps, 1):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(3)
    r1 = p.add_run(f"  {i:02d}.  ")
    set_font(r1, bold=True, color=(31, 73, 125))
    r2 = p.add_run(s)
    set_font(r2, size=10.5)

# ══════════════════════════════════════════════════════════════════════════════
#  CONTACT & DISCLAIMER
# ══════════════════════════════════════════════════════════════════════════════
doc.add_paragraph()
divider()
doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("CONTACT US")
set_font(r, bold=True, size=12, color=(31, 73, 125))

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Shiftzy Go Technologies Pvt. Ltd.\n"
              "Email: business@shiftzygo.com  |  Phone: [Your Number]\n"
              "Website: www.shiftzygo.com")
set_font(r, size=10, color=(89, 89, 89))

doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run(
    "This document is strictly confidential and intended solely for the named recipient. "
    "It does not constitute a legally binding agreement. All financial figures are forward-looking "
    "projections subject to change. This proposal does not constitute insurance or legal advice — "
    "please consult qualified professionals before acting on any information herein."
)
set_font(r, size=9, italic=True, color=(128, 128, 128))

# ── Save ──────────────────────────────────────────────────────────────────────
out_path = "Shiftzy_Go_Insurance_Partnership_Proposal.docx"
doc.save(out_path)
print(f"Saved: {out_path}")
